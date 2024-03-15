from docx import Document
from openpyxl import load_workbook
import tkinter as tk
from tkinter import filedialog, messagebox
import os

def replace_text():
    #Выбираем файл Xlsx с помощью диалогового окна
    xlsx_path = filedialog.askopenfilename(
        initialdir=os.getcwd(),
        title="Выберите данные файл с данными",
        filetypes=(("Excel files", "*.xlsx"), ("all files", "*.*"))
    )
    #Выбираем файл Docx с помощью диалогового окна
    docx_path = filedialog.askopenfilename(
        initialdir=os.getcwd(),
        title="Выберите файл шаблона для заполнения",
        filetypes=(("Word files", "*.docx"), ("all files", "*.*"))
    )
    #берем название листа из названия файла и используем для заполнения словаря для замены
    sheet_name = os.path.basename(docx_path).split('.')[0]
    wb = load_workbook(xlsx_path)
    try:
        ws = wb[sheet_name]
    except KeyError:
        messagebox.showerror("Ой, ошибочка", "Лист с данными не найден. \n"
                              "Выберите файл docx анлогичный названию листа в xlsx")
    data = {}
    for row in ws.iter_rows():
        if row[0].value is not None:
            data[row[0].value] = row[1].value

    #Сохраняем название и убираем расширение из названия файла
    old_doc_name, ext = os.path.splitext(os.path.basename(docx_path))
    assert ext == '.docx'
    #открыть документ шаблона
    doc = Document(docx_path)

    #заменить текст в документе
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.strike:
                for key, value in data.items():
                    if key in run.text:
                            run.font.strike = False
                            run.text = run.text.replace(key, value)
    #заменить текст в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.strike:
                            for key, value in data.items():
                                  if key in run.text:
                                    run.font.strike = False
                                    run.text = run.text.replace(key, value)

    #сохранить документ с новым именем
    if os.path.isfile(old_doc_name + '_редактирован.docx'):
        os.remove(old_doc_name + '_редактирован.docx')
    try:
        doc.save(os.path.join(os.path.dirname(docx_path), os.path.basename(old_doc_name + '_редактирован.docx')))
        messagebox.showinfo("Ой, получилось?", "Файл сохранен")
    except Exception as e:
        messagebox.showerror("Ой, ошибочка", "Закройте документ _редактирован. \n"
                              "error: " + str(e))


def test_replace():
    docx_path = filedialog.askopenfilename(
        initialdir=os.getcwd(),
        title="Выберите файл шаблона для заполнения",
        filetypes=(("Word files", "*.docx"), ("all files", "*.*"))
    )
    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.strike:
                messagebox.showinfo("Проверь эти слова", run.text)


def help():
    messagebox.showinfo(
        "Help",
        "Привет, я заменяю зачеркнутые слова в word.\n"
        "Чтобы составить базу слов для замены, заполни таблицу в xlsx по шаблону Data.xlsx.\n"
        "Главное чтобы название шаблона docx совпадало с названием листа в xlsx. \n" 
        "Если какой то текст не заменяется, попробуйте поместить его в таблицу для надежности. \n"
        "АХО help. \n"
        "Кнопка Найти ошибки существует только для поиска зачеркнутых слов которые не заменились автоматом в _редактирован."
    )


root = tk.Tk()
root.iconbitmap("/Users/leo/Documents/autotext/icon.ico")
root.title("Change Strike text In Word")
root.geometry("332x213")
background_image = tk.PhotoImage(file="/Users/leo/Documents/autotext/background.png")
background_label = tk.Label(root, image=background_image)
background_label.place(relwidth=1, relheight=1)
button = tk.Button(root, text="Заменить текст в документе", command=replace_text, height=2)
help_button = tk.Button(root, text="Помощь", command=help, height=1)
test_button = tk.Button(root, text="Найти ошибки", command=test_replace, height=1)
button.pack(side="bottom", fill="x")
help_button.pack(side="top", fill="x")
test_button.pack(side="top", fill="x")
root.mainloop()


